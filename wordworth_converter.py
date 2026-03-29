#!/usr/bin/env python3
"""
Amiga Wordworth Converter
Supports IFF WORD (Wordworth 2/3) and IFF WOWO (Wordworth 4+) formats.
Converts to .docx and/or .txt

Requires:  python-docx   (pip install python-docx)
Optional:  tkinterdnd2   (pip install tkinterdnd2)  <- true drag-and-drop
"""

import struct
import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox
from pathlib import Path


# =============================================================
#  IFF / Wordworth Parser
# =============================================================
#
#  WORD format (Wordworth 2/3) — chunk sequence:
#    PARA [16 bytes]  <- paragraph group properties (alignment etc.)
#    TABS [variable]  <- tab stop list (optional)
#    TEXT [variable]  <- one visual line, Latin-1 encoded
#    FSCC [8 bytes]   <- inline font/style change (between runs)
#    TEXT ...         <- more lines sharing same PARA properties
#    PARA ...         <- new paragraph group
#    HEAD / FOOT      <- header/footer region markers (skip contents)
#
#  WOWO format (Wordworth 4+) — same IFF container, W-prefixed chunks:
#    WPAR [36 bytes]  <- paragraph properties
#    WTAB [variable]  <- tab stop list (optional)
#    WTXT [variable]  <- one paragraph; 0x0F bytes mark visual line breaks
#    WFSC [variable]  <- inline font/style change
#    WHED / WFOT      <- header/footer region markers (skip contents)
#
#  In WOWO, each WTXT is a whole paragraph (word-wrapped prose has 0x0F
#  as line-break markers within it). We join segments back into one line.


def is_wordworth_file(path):
    """Quick magic-byte check. Returns 'WORD', 'WOWO', or None."""
    try:
        with open(path, 'rb') as f:
            header = f.read(12)
        if len(header) >= 12 and header[:4] == b'FORM':
            ft = header[8:12]
            if ft in (b'WORD', b'WOWO'):
                return ft.decode('latin-1')
    except OSError:
        pass
    return None


class ParaProps:
    """Paragraph properties from a PARA chunk (WORD format, 16 bytes)."""
    ALIGN_MAP = {0: 'left', 1: 'right', 2: 'center', 3: 'justify'}

    def __init__(self, data: bytes):
        self.alignment = 'left'
        if len(data) >= 8:
            self.alignment = self.ALIGN_MAP.get(data[7] & 0x03, 'left')


class WowoParaProps:
    """Paragraph properties from a WPAR chunk (WOWO format, 36 bytes).
    Alignment byte position is not yet confirmed; defaults to left."""

    def __init__(self, _data: bytes):
        self.alignment = 'left'


class Line:
    """One visual line / paragraph from a Wordworth document."""

    def __init__(self, raw: bytes, props):
        self.props = props
        self.alignment = props.alignment
        # Decode Latin-1 (Amiga's default encoding), strip nulls
        self.text = raw.decode('latin-1', errors='replace').replace('\x00', '')
        # Map Amiga middle-dot bullet (0xB7) to Unicode bullet
        self.text = self.text.replace('\xb7', '\u2022')

    @property
    def is_empty(self):
        return not self.text.strip()

    @property
    def is_bullet(self):
        return self.text.lstrip('\t ').startswith('\u2022')

    @property
    def bullet_text(self):
        t = self.text.lstrip('\t ')
        return t[1:].lstrip() if t.startswith('\u2022') else t

    @property
    def looks_like_heading(self):
        t = self.text.strip()
        return bool(t) and not self.is_bullet and len(t) <= 70 and t.endswith(':')

    @property
    def is_separator(self):
        """A line of underscores used as a visual rule."""
        t = self.text.strip()
        return len(t) >= 10 and all(c == '_' for c in t)

    @property
    def indent_level(self):
        """Number of leading tab characters."""
        return len(self.text) - len(self.text.lstrip('\t'))

    @property
    def clean(self):
        """Text with tab columns collapsed to triple-space for plain text."""
        return re.sub(r'\t+', '   ', self.text).rstrip()


class WordworthFile:
    """Parses an Amiga Wordworth IFF file (WORD or WOWO format)."""

    def __init__(self, path):
        self.path = path
        self.lines = []
        self.format_type = None
        self._parse()

    def _parse(self):
        with open(self.path, 'rb') as f:
            data = f.read()
        if len(data) < 12:
            raise ValueError("File too small to be an IFF file.")
        if data[:4] != b'FORM':
            raise ValueError("Not an IFF FORM file.")
        form_size = struct.unpack_from('>I', data, 4)[0]
        form_type = data[8:12]
        if form_type == b'WORD':
            self.format_type = 'WORD'
            self._walk_word(data, 12, 8 + form_size)
        elif form_type == b'WOWO':
            self.format_type = 'WOWO'
            self._walk_wowo(data, 12, 8 + form_size)
        else:
            raise ValueError(f"Not a Wordworth file (FORM type {form_type!r}).")

    def _walk_word(self, data, offset, end):
        """Walk IFF WORD chunks (Wordworth 2/3 format)."""
        cur_props = ParaProps(b'')
        in_hf = False

        while offset + 8 <= end:
            cid   = data[offset:offset+4].decode('latin-1')
            size  = struct.unpack_from('>I', data, offset+4)[0]
            cdata = data[offset+8: offset+8+size]
            offset += 8 + size + (size % 2)

            if cid in ('HEAD', 'FOOT'):
                in_hf = True
                continue
            if cid in ('INFO', 'DOC ', 'SECT'):
                in_hf = False
                continue
            if in_hf:
                continue

            if cid == 'PARA':
                cur_props = ParaProps(cdata)
            elif cid == 'TEXT' and size > 0:
                self.lines.append(Line(cdata, cur_props))
            # TABS, FSCC, COLR, FONT, PSET, PREC, empty TEXT: ignored

    def _walk_wowo(self, data, offset, end):
        """Walk IFF WOWO chunks (Wordworth 4+ format)."""
        cur_props = WowoParaProps(b'')
        in_hf = False

        while offset + 8 <= end:
            cid   = data[offset:offset+4].decode('latin-1')
            size  = struct.unpack_from('>I', data, offset+4)[0]
            cdata = data[offset+8: offset+8+size]
            offset += 8 + size + (size % 2)

            if cid in ('WHED', 'WFOT'):
                in_hf = True
                continue
            if in_hf:
                continue

            if cid == 'WPAR':
                cur_props = WowoParaProps(cdata)
            elif cid == 'WTXT' and size > 0:
                # 0x0F bytes mark visual line breaks within the paragraph
                # (where Wordworth word-wrapped on screen). Strip leading
                # spaces from each segment and join into one paragraph.
                segments = cdata.split(b'\x0f')
                parts = [s.lstrip(b' ') for s in segments if s.strip()]
                cleaned = b' '.join(parts)
                self.lines.append(Line(cleaned, cur_props))
            # WTAB, WFSC, WSPC, WBUL, WSTY, WFNT, WINF, WDOC,
            # embedded FORM (GTID index), etc.: all ignored


# =============================================================
#  Plain-text renderer
# =============================================================

def render_txt(ww: WordworthFile) -> str:
    out = []
    prev_empty = False

    for line in ww.lines:
        if line.is_empty:
            if not prev_empty:   # collapse consecutive blanks to one
                out.append('')
            prev_empty = True
            continue
        prev_empty = False

        if line.is_separator:
            out.append('_' * 60)
        elif line.is_bullet:
            indent = '  ' * max(1, line.indent_level)
            out.append(f'{indent}\u2022 {line.bullet_text.strip()}')
        else:
            out.append(line.clean)

    return '\n'.join(out).rstrip() + '\n'


# =============================================================
#  DOCX renderer
# =============================================================

def render_docx(ww: WordworthFile, out_path: str):
    try:
        from docx import Document as DocxDoc
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError(
            "python-docx is not installed.\n"
            "Install with:  pip install python-docx"
        )

    doc = DocxDoc()

    # US Letter, ~0.9" margins
    sec = doc.sections[0]
    sec.page_width  = int(8.5 * 914400)
    sec.page_height = int(11  * 914400)
    m = int(0.9 * 914400)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = m

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    ALIGN = {
        'left':    WD_ALIGN_PARAGRAPH.LEFT,
        'right':   WD_ALIGN_PARAGRAPH.RIGHT,
        'center':  WD_ALIGN_PARAGRAPH.CENTER,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    def bottom_border(p):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '000000')
        pBdr.append(bot)
        pPr.append(pBdr)

    def sp(p, before=0, after=4):
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)

    def add_run(p, text, bold=False, italic=False, size=11):
        r = p.add_run(text)
        r.font.name  = 'Arial'
        r.font.size  = Pt(size)
        r.bold       = bold
        r.italic     = italic
        return r

    first_real = True    # track first non-empty line for name detection
    prev_empty = False

    for line in ww.lines:

        # ── blank line ──────────────────────────────────────
        if line.is_empty:
            if not prev_empty:
                sp(doc.add_paragraph(''), 0, 0)
            prev_empty = True
            continue
        prev_empty = False

        # ── separator rule ─────────────────────────────────
        if line.is_separator:
            p = doc.add_paragraph('')
            bottom_border(p)
            sp(p, 4, 4)
            continue

        # ── section heading ─────────────────────────────────
        if line.looks_like_heading:
            p = doc.add_paragraph('')
            add_run(p, line.text.strip(), bold=True, size=12)
            bottom_border(p)
            sp(p, 10 if not first_real else 0, 4)
            first_real = False
            continue

        # ── bullet ──────────────────────────────────────────
        if line.is_bullet:
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            lvl = line.indent_level
            p.paragraph_format.left_indent       = Inches(0.3 + lvl * 0.25)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            add_run(p, line.bullet_text.strip())
            sp(p, 0, 2)
            first_real = False
            continue

        # ── normal line ─────────────────────────────────────
        p = doc.add_paragraph('')
        p.alignment = ALIGN.get(line.alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # Detect name line: first real content, short, no colon/bullet
        is_name = (first_real
                   and len(line.text.strip()) < 50
                   and ':' not in line.text
                   and not line.is_bullet)
        if is_name:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, line.text.strip(), bold=True, size=16)
        else:
            collapsed = re.sub(r'\t+', '   ', line.text).rstrip()
            add_run(p, collapsed)

        sp(p, 0, 4)
        first_real = False

    doc.save(out_path)


# =============================================================
#  GUI
# =============================================================

class App(tk.Tk):

    ACCENT  = '#5B4FCF'
    BG      = '#1e1e2e'
    PANEL   = '#2a2a3e'
    FG      = '#cdd6f4'
    FG_DIM  = '#888aaa'
    SUCCESS = '#a6e3a1'
    ERROR   = '#f38ba8'
    WARNING = '#fab387'

    def __init__(self):
        super().__init__()
        self.title('Amiga Wordworth Converter')
        self.resizable(True, True)
        self.minsize(620, 520)
        self.configure(bg=self.BG)

        self.var_docx      = tk.BooleanVar(value=True)
        self.var_txt       = tk.BooleanVar(value=True)
        self.var_outdir    = tk.StringVar(value='Same folder as input file')
        self.custom_outdir = None
        # Queue items: (path: str, base: str|None)
        # base is the scanned folder root; None for individually added files.
        self._queue        = []

        self._build_ui()
        self._setup_dnd()

    # ── build UI ─────────────────────────────────────────────

    def _build_ui(self):
        tk.Frame(self, bg=self.ACCENT, height=4).pack(fill='x')

        hdr = tk.Frame(self, bg=self.BG, pady=12)
        hdr.pack(fill='x', padx=20)
        tk.Label(hdr, text='Amiga Wordworth Converter',
                 bg=self.BG, fg=self.FG,
                 font=('Arial', 16, 'bold')).pack(side='left')
        tk.Label(hdr, text='IFF WORD / WOWO \u2192 .docx / .txt',
                 bg=self.BG, fg=self.FG_DIM,
                 font=('Arial', 10)).pack(side='left', padx=12)

        dz = tk.Frame(self, bg=self.PANEL,
                      highlightbackground=self.ACCENT,
                      highlightthickness=2)
        dz.pack(fill='x', padx=20, pady=(0, 10))
        self.drop_label = tk.Label(
            dz,
            text='Drop Wordworth files here\nor click to browse',
            bg=self.PANEL, fg=self.FG_DIM,
            font=('Arial', 12), pady=28, cursor='hand2'
        )
        self.drop_label.pack(fill='x')
        self.drop_label.bind('<Button-1>', lambda e: self._browse())

        opt = tk.Frame(self, bg=self.BG)
        opt.pack(fill='x', padx=20, pady=4)
        tk.Label(opt, text='Output:', bg=self.BG, fg=self.FG,
                 font=('Arial', 10)).pack(side='left')
        for var, lbl in [(self.var_docx, '.docx  (Word)'),
                         (self.var_txt,  '.txt   (plain text)')]:
            tk.Checkbutton(opt, text=lbl, variable=var,
                           bg=self.BG, fg=self.FG, selectcolor=self.PANEL,
                           activebackground=self.BG, activeforeground=self.FG,
                           font=('Arial', 10)).pack(side='left', padx=8)

        df = tk.Frame(self, bg=self.BG)
        df.pack(fill='x', padx=20, pady=4)
        tk.Label(df, text='Save to:', bg=self.BG, fg=self.FG,
                 font=('Arial', 10)).pack(side='left')
        tk.Label(df, textvariable=self.var_outdir,
                 bg=self.BG, fg=self.FG_DIM,
                 font=('Arial', 10)).pack(side='left', padx=8)
        tk.Button(df, text='Choose\u2026', bg=self.PANEL, fg=self.FG,
                  relief='flat', padx=8, pady=2, font=('Arial', 9),
                  command=self._choose_outdir).pack(side='left', padx=4)
        tk.Button(df, text='Reset', bg=self.PANEL, fg=self.FG_DIM,
                  relief='flat', padx=8, pady=2, font=('Arial', 9),
                  command=self._reset_outdir).pack(side='left', padx=2)

        bf = tk.Frame(self, bg=self.BG)
        bf.pack(fill='x', padx=20, pady=8)
        self.convert_btn = tk.Button(
            bf, text='Convert Files',
            bg=self.ACCENT, fg='white', relief='flat',
            padx=16, pady=8, font=('Arial', 11, 'bold'),
            command=self._convert_queued, state='disabled'
        )
        self.convert_btn.pack(side='left')
        tk.Button(bf, text='Clear Queue', bg=self.PANEL, fg=self.FG_DIM,
                  relief='flat', padx=12, pady=8, font=('Arial', 10),
                  command=self._clear).pack(side='left', padx=8)
        tk.Button(bf, text='Scan Folder\u2026', bg=self.PANEL, fg=self.FG,
                  relief='flat', padx=12, pady=8, font=('Arial', 10),
                  command=self._scan_folder).pack(side='left')

        lf = tk.Frame(self, bg=self.BG)
        lf.pack(fill='both', expand=True, padx=20, pady=(4, 0))
        tk.Label(lf, text='Queue:', bg=self.BG, fg=self.FG_DIM,
                 font=('Arial', 9)).pack(anchor='w')
        sb = tk.Scrollbar(lf)
        sb.pack(side='right', fill='y')
        self.listbox = tk.Listbox(
            lf, bg=self.PANEL, fg=self.FG,
            selectbackground=self.ACCENT,
            font=('Courier', 10), relief='flat', bd=0,
            height=7, yscrollcommand=sb.set
        )
        self.listbox.pack(fill='both', expand=True)
        sb.config(command=self.listbox.yview)

        lgf = tk.Frame(self, bg=self.BG)
        lgf.pack(fill='both', expand=True, padx=20, pady=(8, 16))
        tk.Label(lgf, text='Log:', bg=self.BG, fg=self.FG_DIM,
                 font=('Arial', 9)).pack(anchor='w')
        self.log = tk.Text(
            lgf, bg=self.PANEL, fg=self.FG,
            font=('Courier', 9), relief='flat', bd=0,
            height=6, state='disabled', wrap='word'
        )
        self.log.pack(fill='both', expand=True)
        for tag, color in [('ok', self.SUCCESS), ('err', self.ERROR),
                            ('warn', self.WARNING), ('info', self.FG_DIM)]:
            self.log.tag_config(tag, foreground=color)

    def _setup_dnd(self):
        try:
            from tkinterdnd2 import DND_FILES
            self.drop_label.drop_target_register(DND_FILES)
            self.drop_label.dnd_bind('<<Drop>>', self._on_drop)
        except Exception:
            self.drop_label.configure(
                text='Click to browse for Wordworth files\n'
                     '(pip install tkinterdnd2 for drag-and-drop)'
            )

    # ── queue ─────────────────────────────────────────────────

    def _add_files(self, paths, base=None):
        """Add paths to the queue. base is the scan root for folder-mode."""
        added = 0
        existing = {p for p, _ in self._queue}
        for p in paths:
            p = p.strip().strip('{}')
            if not p or not os.path.isfile(p) or p in existing:
                continue
            self._queue.append((p, base))
            existing.add(p)
            label = os.path.relpath(p, base) if base else os.path.basename(p)
            self.listbox.insert('end', label)
            added += 1
        if added:
            self.convert_btn.configure(state='normal')
            self._log(f'Added {added} file(s).', 'info')

    def _clear(self):
        self._queue.clear()
        self.listbox.delete(0, 'end')
        self.convert_btn.configure(state='disabled')

    # ── events ────────────────────────────────────────────────

    def _on_drop(self, event):
        self._add_files(re.findall(r'\{[^}]+\}|\S+', event.data))

    def _browse(self):
        paths = filedialog.askopenfilenames(
            title='Select Wordworth files',
            filetypes=[('All files', '*.*'),
                       ('Wordworth files', '*.wwp *.WW *.txt')]
        )
        if paths:
            self._add_files(list(paths))

    def _scan_folder(self):
        folder = filedialog.askdirectory(title='Select folder to scan for Wordworth files')
        if not folder:
            return
        self._log(f'Scanning {folder} \u2026', 'info')
        found = []
        counts = {}
        for root, dirs, files in os.walk(folder):
            dirs[:] = sorted(d for d in dirs if not d.startswith('.'))
            for fname in sorted(files):
                fpath = os.path.join(root, fname)
                fmt = is_wordworth_file(fpath)
                if fmt:
                    found.append(fpath)
                    counts[fmt] = counts.get(fmt, 0) + 1
        if found:
            self._add_files(found, base=folder)
            parts = [f'{v} {k}' for k, v in sorted(counts.items()) if v]
            self._log(f'Found {len(found)} file(s): {", ".join(parts)}.', 'ok')
        else:
            self._log('No Wordworth files found in that folder.', 'warn')

    def _choose_outdir(self):
        d = filedialog.askdirectory(title='Choose output directory')
        if d:
            self.custom_outdir = d
            self.var_outdir.set(d)

    def _reset_outdir(self):
        self.custom_outdir = None
        self.var_outdir.set('Same folder as input file')

    # ── conversion ────────────────────────────────────────────

    def _convert_queued(self):
        if not self._queue:
            return
        if not self.var_docx.get() and not self.var_txt.get():
            messagebox.showwarning('No output format',
                                   'Please tick at least one output format.')
            return
        ok = err = 0
        for path, base in self._queue:
            try:
                self._convert_one(path, base)
                ok += 1
            except Exception as exc:
                self._log(f'\u2717 {os.path.basename(path)}: {exc}', 'err')
                err += 1
        tag = 'ok' if err == 0 else ('warn' if ok > 0 else 'err')
        self._log(f'Done \u2014 {ok} succeeded, {err} failed.', tag)

    def _convert_one(self, path, base=None):
        self._log(f'Converting {os.path.basename(path)} \u2026', 'info')
        ww   = WordworthFile(path)
        stem = Path(path).stem

        if self.custom_outdir and base:
            # Maintain folder structure relative to the scan root
            rel_dir = os.path.dirname(os.path.relpath(path, base))
            out_dir = os.path.join(self.custom_outdir, rel_dir)
            os.makedirs(out_dir, exist_ok=True)
        else:
            out_dir = self.custom_outdir or str(Path(path).parent)

        if self.var_txt.get():
            p = os.path.join(out_dir, stem + '.txt')
            # Avoid overwriting the input when it already has a .txt extension
            if os.path.normcase(os.path.abspath(p)) == os.path.normcase(os.path.abspath(path)):
                p = os.path.join(out_dir, stem + '_plain.txt')
            with open(p, 'w', encoding='utf-8') as f:
                f.write(render_txt(ww))
            self._log(f'  \u2713 {os.path.basename(p)}', 'ok')

        if self.var_docx.get():
            p = os.path.join(out_dir, stem + '.docx')
            render_docx(ww, p)
            self._log(f'  \u2713 {os.path.basename(p)}', 'ok')

    def _log(self, msg, tag='info'):
        self.log.configure(state='normal')
        self.log.insert('end', msg + '\n', tag)
        self.log.see('end')
        self.log.configure(state='disabled')


# =============================================================
#  Entry point
# =============================================================

def main():
    try:
        from tkinterdnd2 import TkinterDnD
        App.__bases__ = (TkinterDnD.Tk,)
    except ImportError:
        pass
    App().mainloop()


if __name__ == '__main__':
    main()
